import argparse
import os
import pathlib
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
from os.path import abspath, exists, join, splitext

from .version import __version__
from docx import Document


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_MAP = {"w": WORD_NS}
INS_BEGIN, INS_END, DEL_BEGIN, DEL_END = "<ins>", "</ins>", "<del>", "</del>"
NWORDS_AROUND = 4


from docx.oxml.ns import qn

class ChangeDetector:
    def __init__(self, n_words_around=NWORDS_AROUND, output_format='diff'):
        self.n_words_around = n_words_around
        self.output_format = output_format

    def get_text(self, element):
        """Extracts text from w:t and w:delText elements within the given element."""
        # Note: w:delText is used within w:del, w:t is used elsewhere.
        text_parts = []
        target_tags = (qn('w:t'), qn('w:delText'))
        for node in element.iterdescendants():
            if node.tag in target_tags:
                if node.text:
                    text_parts.append(node.text)
        return "".join(text_parts)

    def _get_context_words(self, words, count, from_end=False):
        if from_end:
            return words[-count:] if count > 0 else []
        return words[:count] if count > 0 else []

    def process_paragraph(self, p_element):
        """Processes a w:p element and returns a list of detected changes with context."""
        # We iterate through the children of w:p.
        # Children can be w:pPr, w:r, w:ins, w:del, etc.
        
        changes = []
        # List of (type, content) where type is 'text', 'ins', or 'del'
        segments = []
        
        for child in p_element:
            tag = child.tag
            if tag == qn('w:pPr'):
                if child.find(qn('w:pPrChange')) is not None:
                    segments.append(('fmt', '[para fmt]'))
                continue
            
            if tag == qn('w:ins'):
                content = self.get_text(child)
                if content:
                    segments.append(('ins', content))
            elif tag == qn('w:del'):
                content = self.get_text(child)
                if content:
                    segments.append(('del', content))
            elif tag == qn('w:r'):
                # Check for rPrChange
                rPr = child.find(qn('w:rPr'))
                if rPr is not None and rPr.find(qn('w:rPrChange')) is not None:
                    segments.append(('fmt', '[fmt]'))
                
                content = self.get_text(child)
                if content:
                    segments.append(('text', content))
            else:
                # Assume it's a regular run or other element containing text
                content = self.get_text(child)
                if content:
                    segments.append(('text', content))
        
        # Merge consecutive text segments to ensure proper context extraction
        merged_segments = []
        for seg_type, content in segments:
            if merged_segments and merged_segments[-1][0] == 'text' and seg_type == 'text':
                merged_segments[-1] = ('text', merged_segments[-1][1] + content)
            else:
                merged_segments.append((seg_type, content))
        segments = merged_segments

        if not any(s[0] in ('ins', 'del', 'fmt') for s in segments):
            return []

        # Now group segments and add context
        return self._group_segments(segments)

    def _group_segments(self, segments):
        processed_changes = []
        
        i = 0
        while i < len(segments):
            if segments[i][0] == 'text':
                i += 1
                continue
            
            # Found a change. Group consecutive or near changes.
            start_idx = i
            end_idx = i
            
            # Look ahead to group changes
            while end_idx + 1 < len(segments):
                next_seg = segments[end_idx + 1]
                if next_seg[0] != 'text':
                    end_idx += 1
                    continue
                
                # If it's text, check if it's short enough to keep grouping
                text_content = next_seg[1]
                if len(text_content.split()) <= self.n_words_around:
                    # Check if there is another change after this text
                    found_next_change = False
                    for j in range(end_idx + 2, len(segments)):
                        if segments[j][0] != 'text':
                            found_next_change = True
                            break
                        if len(segments[j][1].split()) > self.n_words_around:
                            break
                    
                    if found_next_change:
                        end_idx += 1
                        continue
                
                break
            
            # Now build the change string with context
            # Left context
            left_context = ""
            if start_idx > 0 and segments[start_idx - 1][0] == 'text':
                words = segments[start_idx - 1][1].split()
                context_words = self._get_context_words(words, self.n_words_around, from_end=True)
                left_context = " ".join(context_words)
                if segments[start_idx - 1][1].endswith(" ") and left_context:
                    left_context += " "
            
            # Right context
            right_context = ""
            if end_idx + 1 < len(segments) and segments[end_idx + 1][0] == 'text':
                words = segments[end_idx + 1][1].split()
                context_words = self._get_context_words(words, self.n_words_around, from_end=False)
                right_context = " ".join(context_words)
                if segments[end_idx + 1][1].startswith(" ") and right_context:
                    right_context = " " + right_context

            if self.output_format == 'tags':
                # Middle content (the changes)
                change_content = ""
                for j in range(start_idx, end_idx + 1):
                    seg_type, content = segments[j]
                    if seg_type == 'ins':
                        change_content += INS_BEGIN + content + INS_END
                    elif seg_type == 'del':
                        change_content += DEL_BEGIN + content + DEL_END
                    else:
                        change_content += content
                processed_changes.append(left_context + change_content + right_context)
            else: # diff format
                prev_middle = ""
                after_middle = ""
                for j in range(start_idx, end_idx + 1):
                    seg_type, content = segments[j]
                    if seg_type == 'ins':
                        after_middle += content
                    elif seg_type == 'del':
                        prev_middle += content
                    else:
                        prev_middle += content
                        after_middle += content
                
                change_str = f"{left_context}{prev_middle}{right_context} -> {left_context}{after_middle}{right_context}"
                processed_changes.append(change_str)

            i = end_idx + 1
            
        return processed_changes



class DocxReviews:
    def __init__(self, file_docx, output_format='diff') -> None:
        assert exists(file_docx)
        self.reviews = []
        self.file_docx = abspath(file_docx)
        self.output_format = output_format
        # use tmp file
        self.target_file = join(tempfile.gettempdir(), "docx_reviews_to_txt.docx")
        if exists(self.target_file):
            os.remove(self.target_file)
            assert not exists(self.target_file)
        try:
            shutil.copyfile(file_docx, self.target_file)
        except Exception as exc:
            # at windows, shutil.copy fail if docx opened and only can be copied from powershell
            if os.name == "nt":
                cmd = f"Copy-Item {file_docx} {self.target_file}"
                subprocess.run(
                    ["powershell", "-noprofile", "-Command", cmd], capture_output=True, check=True
                )
            else:
                raise exc
        assert exists(self.target_file)
        self.paragraphs = Document(self.target_file).paragraphs

    def _extract_comments(self, doc) -> list[str]:
        comments = []
        try:
            # Find the comments part
            for rel in doc.part.rels.values():
                if "comments" in rel.target_ref:
                    root = etree.fromstring(rel.target_part.blob)
                    for comment in root.xpath("//w:comment", namespaces=NS_MAP):
                        author = comment.get(qn("w:author"), "Unknown")
                        text_parts = []
                        for t in comment.xpath(".//w:t", namespaces=NS_MAP):
                            if t.text:
                                text_parts.append(t.text)
                        text = "".join(text_parts)
                        if text:
                            comments.append(f"Comment by {author}: {text}")
        except Exception:
            pass
        return comments

    def _parse(self) -> None:
        detector = ChangeDetector(output_format=self.output_format)
        
        doc = Document(self.target_file)
        for p in doc.paragraphs:
            p_changes = detector.process_paragraph(p._p)
            for change in p_changes:
                self.reviews.append("- " + change)
        
        # Also process tables
        for table in doc.tables:
            for row in table.rows:
                # Check for row-level changes
                tr = row._tr
                trPr = tr.find(qn('w:trPr'))
                row_change = ""
                if trPr is not None:
                    if trPr.find(qn('w:ins')) is not None:
                        row_change = "[Row Ins] "
                    elif trPr.find(qn('w:del')) is not None:
                        row_change = "[Row Del] "

                for cell in row.cells:
                    # Check for cell-level changes
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    cell_change = ""
                    if tcPr is not None:
                        if tcPr.find(qn('w:ins')) is not None:
                            cell_change = "[Cell Ins] "
                        elif tcPr.find(qn('w:del')) is not None:
                            cell_change = "[Cell Del] "

                    for p in cell.paragraphs:
                        p_changes = detector.process_paragraph(p._p)
                        if p_changes:
                            for change in p_changes:
                                self.reviews.append("- " + row_change + cell_change + change)
                        elif row_change or cell_change:
                            # If row/cell changed but no paragraph changes, still report content
                            content = detector.get_text(p._p)
                            if content:
                                self.reviews.append("- " + row_change + cell_change + content)

        # Process comments
        comments = self._extract_comments(doc)
        if comments:
            self.reviews.append("\nComments:")
            for comment in comments:
                self.reviews.append("- " + comment)

    def save_reviews(self) -> None:
        if not self.reviews:
            self._parse()
        filename = splitext(self.file_docx)[0] + "_review.txt"
        with open(filename, "w") as file:
            for change in self.reviews:
                file.write(f"{change}\n")
        assert filename
        print(f"txt reviews at {pathlib.Path(filename).as_uri()}")

    def save_xml_p_elems(self) -> None:
        filename = splitext(self.file_docx)[0] + ".xml"
        with open(filename, "w") as file:
            for p in self.paragraphs:
                xml = p._p.xml
                file.write(f"{xml}\n")
        assert filename
        print(f"xml paragraphs at {pathlib.Path(filename).as_uri()}")


def docxreviews_cli(argv=None) -> None:
    parser = argparse.ArgumentParser(
        prog="docxreviews2txt",
        description="Command line tool to extract review changes from a docx file as plain text using HTML tags <ins> and <del>.",
    )
    parser.add_argument("docx", help="input docx", type=pathlib.Path)
    parser.add_argument(
        "--format",
        help="output format: 'diff' (PREVIOUS -> AFTER) or 'tags' (<ins>/<del>). Default is 'diff'.",
        choices=["diff", "tags"],
        default="diff",
    )
    parser.add_argument(
        "--version", help="show version", action="version", version="%(prog)s " + __version__
    )
    args = parser.parse_args(argv)
    docx_reviews = DocxReviews(file_docx=args.docx, output_format=args.format)
    docx_reviews.save_reviews()
