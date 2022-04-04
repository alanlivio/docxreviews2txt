#!/usr/bin/env python
from docx import Document
import xml.etree.ElementTree as ET
import argparse
import os
import sys
import subprocess
import pathlib
from os.path import exists
import zipfile
import tempfile
import shutil


WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_MAP = {"w": WORD_NS}
ET_WORD_NS = '{' + WORD_NS + '}'
ET_TEXT = ET_WORD_NS + 't'
ET_DEL = ET_WORD_NS + 'del'
ET_INS = ET_WORD_NS + 'ins'
MAX_LEN_LEFT = 4
MAX_LEN_RIGHT = 3


def str_deltext_elms(child):
    x = [text.text for text in child.findall(
        './/w:delText', NS_MAP)]
    return "".join(x)


def str_t_elms(child):
    x = [text.text for text in child.findall(
        './/w:t', NS_MAP)]
    return "".join(x)


def str_left_t_elms(root_p, index):
    left_ar = []
    for i in range(index, 0, -1):
        if (root_p[i].tag == ET_INS):
            continue
        left_ar = str_t_elms(root_p[i]).split(" ") + left_ar
        if(len(left_ar) >= MAX_LEN_LEFT):
            left_ar = left_ar[-MAX_LEN_LEFT:]
        break
    return " ".join(left_ar)


def str_right_t_elms(root_p, index):
    right_ar = []
    for i in range(index, len(root_p)):
        if (root_p[i].tag == ET_INS):
            continue
        right_ar = str_t_elms(root_p[i]).split(" ") + right_ar
        if(len(right_ar) >= MAX_LEN_RIGHT):
            right_ar = right_ar[:MAX_LEN_RIGHT]
        break
    return " ".join(right_ar)


class DocxReviews:
    def __init__(self, file_docx, verbose):
        self.verbose = verbose
        self.reviews = []
        self.file_docx = file_docx
        # extract docxZip and paragraphs
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, 'docx_reviews_to_txt.docx')
        if exists(temp_path):
            os.remove(temp_path)
        try:
            shutil.copy(file_docx, temp_path)
        except:
            # at windows, shutil.copy fail if docx opened and only can be copied from powershell
            if os.name == 'nt':
                cmd = f"Copy-Item {file_docx} {temp_path}"
                subprocess.run(["powershell", "-Command", cmd], capture_output=True)
        self.docxZip = zipfile.ZipFile(temp_path, mode="r")
        self.paragraphs = Document(temp_path).paragraphs

    def reviews_append(self, text):
        if not len(text):
            return
        self.reviews.append(text)
        if self.verbose:
            print(self.reviews[-1])

    def parse(self):
        # comments
        if 'word/comments.xml' in [member.filename for member in self.docxZip.infolist()]:
            commentsXML = self.docxZip.read('word/comments.xml')
            # print(commentsXML)
            root = ET.fromstring(commentsXML)
            comments = root.findall('.//w:comment', NS_MAP)
            if len(comments):
                self.reviews_append("# Comments")
                for comment in comments:
                    lines = comment.findall('.//w:r', NS_MAP)
                    for line in lines:
                        self.reviews_append(str_t_elms(line))

        # changes
        self.reviews_append("# Typos and rewriting suggestions")
        for p in self.paragraphs:
            xml = p._p.xml
            root = ET.fromstring(xml)
            if len(root.findall('.//w:del', NS_MAP)) == 0 and len(root.findall('.//w:ins', NS_MAP)) == 0:
                continue
            for index in range(len(root)-1):
                prev = root[index-1]
                cur = root[index]
                next = root[index+1]
                # DEL followed by INS
                if (cur.tag == ET_DEL and next.tag == ET_INS):
                    del_text = str_deltext_elms(cur)
                    ins_text = str_t_elms(next)
                    left_text = str_left_t_elms(root, index-1)
                    right_text = str_right_t_elms(root, index+2)
                    self.reviews_append("- " + left_text + del_text + right_text +
                                        " -> " + left_text + ins_text + right_text)
                # INS alone
                elif (cur.tag == ET_INS and prev.tag != ET_DEL):
                    ins_text = str_t_elms(cur)
                    left_text = str_left_t_elms(root, index-1)
                    right_text = str_right_t_elms(root, index)
                    self.reviews_append("- " + left_text + right_text +
                                        " -> " + left_text + ins_text + right_text)
                # DEL alone
                elif (cur.tag == ET_DEL and next.tag != ET_INS):
                    del_text = str_deltext_elms(cur)
                    left_text = str_left_t_elms(root, index-1)
                    right_text = str_right_t_elms(root, index+1)
                    self.reviews_append("- " + left_text + del_text + right_text +
                                        " -> " + left_text + right_text)

    def save_reviews_to_file(self):
        file_txt_name = str(os.path.splitext(self.file_docx)[0])+'_review.txt'
        with open(file_txt_name, "w") as file:
            for change in self.reviews:
                file.write(f"{change}\n")

    def save_xml_p_elems(self):
        file_txt_name = str(os.path.splitext(self.file_docx)[0])+'.xml'
        with open(file_txt_name, "w") as file:
            for p in self.paragraphs:
                xml = p._p.xml
                file.write(f"{xml}\n")


def main(argv):
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "docx", help="input docx", type=pathlib.Path)
    parser.add_argument(
        '--save_txt', help='save review as txt', action="store_true")
    parser.add_argument(
        '--save_p_xml', help='save extracted paragraphs xml', action="store_true")
    args = parser.parse_args(argv)
    verbose = not args.save_p_xml and not args.save_txt
    if not exists(args.docx):
        print(f'{args.docx} does not exist')
        return 1
    docx_reviews = DocxReviews(args.docx, verbose)
    docx_reviews.parse()
    if args.save_p_xml:
        docx_reviews.save_xml_p_elems()
    if args.save_txt:
        docx_reviews.save_reviews_to_file()
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
